import os
import json
import logging
from datetime import datetime, timedelta, timezone

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from models import Incident, ThreatPerson, ThreatOrganization, IOCIndicator
from database import get_session
from config import BASE_DIR, IOC_THREAT_LEVEL_LABELS, IOC_TYPE_LABELS

logger = logging.getLogger(__name__)

REPORT_DIR = os.path.join(os.path.expanduser('~'), 'Desktop')
IMAGES_DIR = os.path.join(BASE_DIR, 'data', 'images')

# Max image width in the Word document (inches)
MAX_IMAGE_WIDTH = Inches(5.5)

# Ukrainian month names for formatting
MONTHS_UK = {
    1: 'січня', 2: 'лютого', 3: 'березня', 4: 'квітня',
    5: 'травня', 6: 'червня', 7: 'липня', 8: 'серпня',
    9: 'вересня', 10: 'жовтня', 11: 'листопада', 12: 'грудня',
}


def format_date_uk(dt):
    """Format datetime as '14 лютого 2026'."""
    if not dt:
        return ''
    return f"{dt.day} {MONTHS_UK.get(dt.month, '')} {dt.year}"


def _add_images_to_doc(doc, incident):
    """Add images from incident to the Word document."""
    if not incident.images:
        return

    try:
        image_paths = json.loads(incident.images)
    except (json.JSONDecodeError, TypeError):
        return

    added = 0
    for img_path in image_paths:
        if not os.path.isfile(img_path):
            continue
        try:
            doc.add_picture(img_path, width=MAX_IMAGE_WIDTH)
            # Center the image
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            added += 1
        except Exception as e:
            logger.warning(f"Could not add image {img_path}: {e}")

    if added:
        doc.add_paragraph()  # spacer after images


def generate_daily_report(target_date=None):
    """
    Generate a Word report for incidents on target_date.
    If target_date is None, uses today.
    Returns the path to the generated file.
    """
    if target_date is None:
        target_date = datetime.now(timezone.utc).date()
    elif isinstance(target_date, datetime):
        target_date = target_date.date()

    session = get_session()
    try:
        # Query incidents for the target date
        date_start = datetime(target_date.year, target_date.month, target_date.day, tzinfo=timezone.utc)
        date_end = date_start + timedelta(days=1)

        incidents = (
            session.query(Incident)
            .filter(Incident.date >= date_start, Incident.date < date_end)
            .order_by(Incident.date.desc())
            .all()
        )

        # If no incidents for exact date, get all recent (last 7 days)
        if not incidents:
            week_ago = date_start - timedelta(days=7)
            incidents = (
                session.query(Incident)
                .filter(Incident.date >= week_ago, Incident.date < date_end)
                .order_by(Incident.date.desc())
                .all()
            )

        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Title
        title = doc.add_heading(level=0)
        run = title.add_run(f"Звіт про кібератаки проти України")
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 91, 187)  # Ukrainian blue

        # Date subtitle
        date_str = format_date_uk(date_start)
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(date_str)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)

        # Summary
        doc.add_paragraph()
        summary = doc.add_paragraph()
        run = summary.add_run(f"Всього інцидентів у звіті: {len(incidents)}")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph()  # spacer

        if not incidents:
            p = doc.add_paragraph()
            run = p.add_run("Інцидентів за вказану дату не знайдено.")
            run.italic = True
        else:
            for i, inc in enumerate(incidents, 1):
                # Incident header
                heading = doc.add_heading(level=2)
                run = heading.add_run(f"{i}. {inc.title}")
                run.font.size = Pt(14)

                # Ukrainian translation of title
                if inc.title_uk and inc.title_uk != inc.title:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[UA] {inc.title_uk}")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 91, 187)

                # Metadata table
                meta_items = []
                if inc.date:
                    meta_items.append(('Дата', format_date_uk(inc.date)))
                if inc.attack_type:
                    meta_items.append(('Тип атаки', inc.attack_type))
                if inc.target_sector:
                    meta_items.append(('Сектор', inc.target_sector))
                if inc.severity:
                    meta_items.append(('Критичність', inc.severity))
                if inc.threat_actor:
                    meta_items.append(('Загрозливий актор', inc.threat_actor))
                if inc.source:
                    meta_items.append(('Джерело', inc.source))

                if meta_items:
                    table = doc.add_table(rows=len(meta_items), cols=2)
                    table.style = 'Light Grid Accent 1'
                    for row_idx, (label, value) in enumerate(meta_items):
                        cells = table.rows[row_idx].cells
                        cells[0].text = label
                        cells[1].text = value
                        # Bold the label
                        for paragraph in cells[0].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

                doc.add_paragraph()  # spacer

                # Images from article
                _add_images_to_doc(doc, inc)

                # Full text (Ukrainian translation preferred, fallback to description)
                full_text = inc.full_text_uk or inc.full_text
                desc_text = inc.description_uk or inc.description

                if full_text:
                    p = doc.add_paragraph()
                    run = p.add_run("Повний текст статті:")
                    run.bold = True
                    # Split into paragraphs for better formatting
                    for para_text in full_text.split('\n\n'):
                        para_text = para_text.strip()
                        if para_text:
                            doc.add_paragraph(para_text[:5000])
                elif desc_text:
                    p = doc.add_paragraph()
                    run = p.add_run("Опис:")
                    run.bold = True
                    doc.add_paragraph(desc_text[:3000])

                # Source URL
                if inc.source_url:
                    p = doc.add_paragraph()
                    run = p.add_run("Джерело: ")
                    run.bold = True
                    p.add_run(inc.source_url)

                # Separator
                if i < len(incidents):
                    doc.add_paragraph('_' * 60)
                    doc.add_paragraph()

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("--- CyberTracker UA ---")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)

        # Save
        filename = f"CyberTracker_Report_{target_date.strftime('%Y-%m-%d')}.docx"
        filepath = os.path.join(REPORT_DIR, filename)
        doc.save(filepath)

        logger.info(f"Report generated: {filepath} ({len(incidents)} incidents)")
        return filepath

    finally:
        session.close()


def _add_person_photo(doc, person):
    """Add person photo to the document if available."""
    if not person.photo_url:
        return
    photo_path = os.path.join(BASE_DIR, person.photo_url.lstrip('/'))
    if not os.path.isfile(photo_path):
        logger.warning(f"Photo not found: {photo_path}")
        return
    try:
        doc.add_picture(photo_path, width=Inches(1.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        logger.warning(f"Could not add photo for {person.name}: {e}")


def generate_org_report(org_id):
    """
    Generate a Word report for all persons in a given organization.
    Returns the path to the generated file.
    """
    session = get_session()
    try:
        org = session.query(ThreatOrganization).get(org_id)
        if not org:
            raise ValueError(f"Organization with id={org_id} not found")

        members = session.query(ThreatPerson).filter(
            ThreatPerson.organization.ilike(f'%{org.name}%')
        ).order_by(ThreatPerson.name).all()

        doc = Document()

        # Default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # === Title ===
        title = doc.add_heading(level=0)
        run = title.add_run(f"Звіт: {org.name}")
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 91, 187)

        # Date
        now = datetime.now(timezone.utc)
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Звіт згенеровано: {format_date_uk(now)}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()

        # === Organization Info ===
        heading = doc.add_heading(level=1)
        run = heading.add_run("Інформація про організацію")
        run.font.color.rgb = RGBColor(0, 91, 187)

        org_items = [('Назва', org.name)]
        if org.org_type:
            org_items.append(('Тип', org.org_type))
        if org.aliases:
            org_items.append(('Альтернативні назви', org.aliases))
        if org.country:
            org_items.append(('Країна', org.country))
        if org.parent_org:
            org_items.append(('Головна організація', org.parent_org))
        if org.members_count:
            org_items.append(('Відомих членів', str(org.members_count)))

        table = doc.add_table(rows=len(org_items), cols=2)
        table.style = 'Light Grid Accent 1'
        for row_idx, (label, value) in enumerate(org_items):
            cells = table.rows[row_idx].cells
            cells[0].text = label
            cells[1].text = value
            for paragraph in cells[0].paragraphs:
                for r in paragraph.runs:
                    r.bold = True

        doc.add_paragraph()

        if org.description:
            p = doc.add_paragraph()
            run = p.add_run("Опис: ")
            run.bold = True
            doc.add_paragraph(org.description)

        if org.known_operations:
            try:
                operations = json.loads(org.known_operations)
                if operations:
                    p = doc.add_paragraph()
                    run = p.add_run("Відомі операції організації:")
                    run.bold = True
                    for op in operations:
                        doc.add_paragraph(op, style='List Bullet')
            except (json.JSONDecodeError, TypeError):
                pass

        doc.add_paragraph('_' * 60)
        doc.add_paragraph()

        # === Members ===
        heading = doc.add_heading(level=1)
        run = heading.add_run(f"Члени організації ({len(members)})")
        run.font.color.rgb = RGBColor(0, 91, 187)

        if not members:
            p = doc.add_paragraph()
            run = p.add_run("Членів організації не знайдено в базі.")
            run.italic = True
        else:
            for i, person in enumerate(members, 1):
                person_heading = doc.add_heading(level=2)
                run = person_heading.add_run(f"{i}. {person.name}")
                run.font.size = Pt(14)

                _add_person_photo(doc, person)

                meta_items = []
                if person.aliases:
                    meta_items.append(('Позивні / Alias', person.aliases))
                if person.role:
                    meta_items.append(('Роль', person.role))
                if person.organization:
                    meta_items.append(('Організація', person.organization))
                if person.country:
                    meta_items.append(('Країна', person.country))
                if person.status:
                    meta_items.append(('Статус', person.status))
                if person.source_url:
                    meta_items.append(('Джерело', person.source_url))

                if meta_items:
                    table = doc.add_table(rows=len(meta_items), cols=2)
                    table.style = 'Light Grid Accent 1'
                    for row_idx, (label, value) in enumerate(meta_items):
                        cells = table.rows[row_idx].cells
                        cells[0].text = label
                        cells[1].text = value
                        for paragraph in cells[0].paragraphs:
                            for r in paragraph.runs:
                                r.bold = True

                doc.add_paragraph()

                if person.description:
                    p = doc.add_paragraph()
                    run = p.add_run("Досьє:")
                    run.bold = True
                    doc.add_paragraph(person.description)

                if person.operations:
                    try:
                        ops = json.loads(person.operations)
                        if ops:
                            p = doc.add_paragraph()
                            run = p.add_run("Відомі операції:")
                            run.bold = True
                            for op in ops:
                                doc.add_paragraph(op, style='List Bullet')
                    except (json.JSONDecodeError, TypeError):
                        pass

                if i < len(members):
                    doc.add_paragraph('_' * 60)
                    doc.add_paragraph()

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("--- CyberTracker UA ---")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)

        # Save
        safe_name = org.name
        for ch in '/\\<>:|?*\u00ab\u00bb"\'':
            safe_name = safe_name.replace(ch, '')
        safe_name = safe_name.strip()
        filename = f"CyberTracker_Org_{safe_name}_{now.strftime('%Y-%m-%d')}.docx"
        filepath = os.path.join(REPORT_DIR, filename)
        doc.save(filepath)

        logger.info(f"Org report generated: {filepath} ({len(members)} members)")
        return filepath

    finally:
        session.close()


def generate_ioc_report(ioc_type_filter=None, threat_level_filter=None, source_filter=None):
    """
    Generate a Word report of IOC indicators.
    Optional filters: ioc_type, threat_level, source.
    Returns the path to the generated file.
    """
    session = get_session()
    try:
        query = session.query(IOCIndicator)

        if ioc_type_filter:
            query = query.filter(IOCIndicator.ioc_type == ioc_type_filter)
        if threat_level_filter:
            query = query.filter(IOCIndicator.threat_level == threat_level_filter)
        if source_filter:
            query = query.filter(IOCIndicator.source == source_filter)

        all_iocs = query.order_by(IOCIndicator.threat_level, IOCIndicator.last_seen.desc()).all()

        doc = Document()

        # Default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # === Title ===
        title = doc.add_heading(level=0)
        run = title.add_run("Звіт IOC індикаторів компрометації")
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 91, 187)

        # Date
        now = datetime.now(timezone.utc)
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Звіт згенеровано: {format_date_uk(now)}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()

        # === Summary statistics ===
        heading = doc.add_heading(level=1)
        run = heading.add_run("Зведена статистика")
        run.font.color.rgb = RGBColor(0, 91, 187)

        # Count by threat level
        level_order = ['critical', 'high', 'medium', 'low', 'unknown']
        level_counts = {}
        for ioc in all_iocs:
            lvl = ioc.threat_level or 'unknown'
            level_counts[lvl] = level_counts.get(lvl, 0) + 1

        summary_items = [
            ('Всього IOC у звіті', str(len(all_iocs))),
        ]
        for lvl in level_order:
            cnt = level_counts.get(lvl, 0)
            if cnt > 0:
                label = IOC_THREAT_LEVEL_LABELS.get(lvl, lvl)
                summary_items.append((label, str(cnt)))

        # Count by source
        source_counts = {}
        for ioc in all_iocs:
            src = ioc.source or 'Unknown'
            source_counts[src] = source_counts.get(src, 0) + 1
        for src, cnt in sorted(source_counts.items(), key=lambda x: -x[1]):
            summary_items.append((f"Джерело: {src}", str(cnt)))

        table = doc.add_table(rows=len(summary_items), cols=2)
        table.style = 'Light Grid Accent 1'
        for row_idx, (label, value) in enumerate(summary_items):
            cells = table.rows[row_idx].cells
            cells[0].text = label
            cells[1].text = value
            for paragraph in cells[0].paragraphs:
                for r in paragraph.runs:
                    r.bold = True

        doc.add_paragraph()

        # === IOC grouped by type ===
        iocs_by_type = {}
        for ioc in all_iocs:
            t = ioc.ioc_type or 'other'
            if t not in iocs_by_type:
                iocs_by_type[t] = []
            iocs_by_type[t].append(ioc)

        # Type order
        type_order = ['ipv4', 'ipv6', 'domain', 'url', 'hash_md5', 'hash_sha1', 'hash_sha256', 'email', 'cve']
        sorted_types = [t for t in type_order if t in iocs_by_type]
        sorted_types += [t for t in iocs_by_type if t not in type_order]

        MAX_PER_TYPE = 100

        for ioc_type in sorted_types:
            iocs = iocs_by_type[ioc_type]
            type_label = IOC_TYPE_LABELS.get(ioc_type, ioc_type.upper())

            heading = doc.add_heading(level=2)
            run = heading.add_run(f"{type_label} ({len(iocs)})")
            run.font.size = Pt(14)

            # Table header
            display_iocs = iocs[:MAX_PER_TYPE]
            tbl = doc.add_table(rows=1 + len(display_iocs), cols=5)
            tbl.style = 'Light Grid Accent 1'

            # Header row
            headers = ['Значення', 'Джерело', 'Рівень загрози', 'Дата', 'Теги']
            for col_idx, hdr in enumerate(headers):
                cell = tbl.rows[0].cells[col_idx]
                cell.text = hdr
                for paragraph in cell.paragraphs:
                    for r in paragraph.runs:
                        r.bold = True

            # Data rows
            for row_idx, ioc in enumerate(display_iocs, 1):
                cells = tbl.rows[row_idx].cells
                # Truncate long values for readability
                val = ioc.value[:80] + '...' if len(ioc.value) > 80 else ioc.value
                cells[0].text = val
                cells[1].text = ioc.source or ''
                cells[2].text = IOC_THREAT_LEVEL_LABELS.get(ioc.threat_level, ioc.threat_level or '')
                cells[3].text = ioc.last_seen.strftime('%d.%m.%Y') if ioc.last_seen else ''
                tags_str = ioc.tags[:50] + '...' if ioc.tags and len(ioc.tags) > 50 else (ioc.tags or '')
                cells[4].text = tags_str

            if len(iocs) > MAX_PER_TYPE:
                p = doc.add_paragraph()
                run = p.add_run(f"... та ще {len(iocs) - MAX_PER_TYPE} записів")
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)

            doc.add_paragraph()

        if not all_iocs:
            p = doc.add_paragraph()
            run = p.add_run("IOC індикаторів не знайдено за заданими фільтрами.")
            run.italic = True

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("--- CyberTracker UA ---")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)

        # Save
        filename = f"CyberTracker_IOC_{now.strftime('%Y-%m-%d')}.docx"
        filepath = os.path.join(REPORT_DIR, filename)
        doc.save(filepath)

        logger.info(f"IOC report generated: {filepath} ({len(all_iocs)} IOCs)")
        return filepath

    finally:
        session.close()
